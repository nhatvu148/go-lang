# coding: utf-8

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Mm
from docx2pdf import convert

import os
# import sys
# sys.path.append(r"C:\Users\nhatv\Downloads\Report\images")
directory_path = os.getcwd()
file_path = directory_path + r"\report.docx"
file_path_pdf = directory_path + r"\report.pdf"

##
alignment_dict = {"justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
                  "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                  "centre": WD_PARAGRAPH_ALIGNMENT.CENTER,
                  "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
                  "left": WD_PARAGRAPH_ALIGNMENT.LEFT}

orientation_dict = {"portrait": WD_ORIENTATION.PORTRAIT,
                    "landscape": WD_ORIENTATION.LANDSCAPE}
##

##
pic_size = [[Mm(171), Mm(80)], [Mm(171), Mm(46)]]
margin_top = Mm(26.9)
margin_bottom = Mm(31.1)
margin_left = Mm(19)
margin_right = Mm(17.7)
date_time = "2021 年2 月1 日"
company_name = "日本シップヤード株式会社"
num_of_row = 9
num_of_col = 5
col_titles = ["S.No.5215", "S.No.5167", "S.No.5276"]
row_data = [["運航", "○", "○", "―"],
            ["積付", "×", "×", "―"],
            ["センサー（応力）", "○", "○", "―"],
            ["センサー（レインフロー）", "○", "○", "―"],
            ["センサー（加速度）", "○", "○", "―"],
            ["センサー（運動）", "△（要ノイズ除去）", "○", "―"],
            ["波浪レーダー", "未搭載", "○", "搭載予定"],
            ["海象予報値", "△（SN2 準備中）", "△（SN2 準備中）", "―"]]
font_type = "Meiryo UI"
getting_data = ["S.No.5167", "2020 年12 月", "20 ", "Sea-Navi1.0"]
picture_2_group = [["DFP", os.getcwd() + "/images/" + "graph-Stress-1" + ".png"],
                   ["DFS", os.getcwd() + "/images/" + "graph-Stress-2" + ".png"],
                   ["SFP", os.getcwd() + "/images/" + "graph-Stress-29" + ".png"],
                   ["SFS", os.getcwd() + "/images/" + "graph-Stress-30" + ".png"],
                   ["DMP", os.getcwd() + "/images/" + "graph-Stress-31" + ".png"],
                   ["DMS", os.getcwd() + "/images/" + "graph-Stress-32" + ".png"],
                   ["SMP", os.getcwd() + "/images/" + "graph-Stress-33" + ".png"],
                   ["SMS", os.getcwd() + "/images/" + "graph-Stress-34" + ".png"],
                   ["DAP", os.getcwd() + "/images/" + "graph-Stress-35" + ".png"],
                   ["DAS", os.getcwd() + "/images/" + "graph-Stress-36" + ".png"],
                   ["SAP", os.getcwd() + "/images/" + "graph-Stress-37" + ".png"],
                   ["SAS", os.getcwd() + "/images/" + "graph-Stress-38" + ".png"],
                   ["L21", os.getcwd() + "/images/" + "graph-Stress-39" + ".png"],
                   ["L22", os.getcwd() + "/images/" + "graph-Stress-40" + ".png"],
                   ["L23", os.getcwd() + "/images/" + "graph-Stress-3" + ".png"],
                   ["L24", os.getcwd() + "/images/" + "graph-Stress-4" + ".png"],
                   ["L25", os.getcwd() + "/images/" + "graph-Stress-5" + ".png"],
                   ["L41", os.getcwd() + "/images/" + "graph-Stress-6" + ".png"],
                   ["L42", os.getcwd() + "/images/" + "graph-Stress-7" + ".png"],
                   ["L43", os.getcwd() + "/images/" + "graph-Stress-8" + ".png"],
                   ["L44", os.getcwd() + "/images/" + "graph-Stress-9" + ".png"],
                   ["L45", os.getcwd() + "/images/" + "graph-Stress-10" + ".png"],
                   ["L46", os.getcwd() + "/images/" + "graph-Stress-11" + ".png"],
                   ["L47", os.getcwd() + "/images/" + "graph-Stress-12" + ".png"],
                   ["L48", os.getcwd() + "/images/" + "graph-Stress-13" + ".png"],
                   ["L49", os.getcwd() + "/images/" + "graph-Stress-14" + ".png"],
                   ["L410", os.getcwd() + "/images/" + "graph-Stress-15" + ".png"],
                   ["L51", os.getcwd() + "/images/" + "graph-Stress-16" + ".png"],
                   ["L52", os.getcwd() + "/images/" + "graph-Stress-17" + ".png"],
                   ["L53", os.getcwd() + "/images/" + "graph-Stress-18" + ".png"],
                   ["L54", os.getcwd() + "/images/" + "graph-Stress-19" + ".png"],
                   ["L55", os.getcwd() + "/images/" + "graph-Stress-20" + ".png"],
                   ["L56", os.getcwd() + "/images/" + "graph-Stress-21" + ".png"],
                   ["L57", os.getcwd() + "/images/" + "graph-Stress-22" + ".png"]]
##
picture_3_group = [["AFx", os.getcwd() + "/images/" + "graph-Stress-23" + ".png"],
                   ["AFy", os.getcwd() + "/images/" + "graph-Stress-24" + ".png"],
                   ["AFz", os.getcwd() + "/images/" + "graph-Stress-25" + ".png"],
                   ["AAx", os.getcwd() + "/images/" + "graph-Stress-26" + ".png"],
                   ["AAy", os.getcwd() + "/images/" + "graph-Stress-27" + ".png"],
                   ["AAz", os.getcwd() + "/images/" + "graph-Stress-28" + ".png"]]
##
picture_4_group = [["Roll", os.getcwd() + "/images/" + "graph-Gyro-1" + ".png"],
                   ["Pitch", os.getcwd() + "/images/" + "graph-Gyro-2" + ".png"]]
##
picture_5_group = [["有義波高", os.getcwd() + "/images/" + "graph-Wave-1" + ".png"],
                   ["波周期", os.getcwd() + "/images/" + "graph-Wave-2" + ".png"]]
##

####
doc_file = Document()
sections = doc_file.sections
for section in sections:
    section.top_margin = margin_top
    section.bottom_margin = margin_bottom
    section.left_margin = margin_left
    section.right_margin = margin_right
###

def add_picture(document, 
                path = "", 
                align = "justify",
                pic_width = pic_size[1][0],
                pic_height = pic_size[1][1]):
    document.add_picture(path, 
                         width = pic_width,
                         height = pic_height)
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = alignment_dict[align.lower()]

def add_content(document,
                content, 
                space_after = 0, 
                font_name = font_type, 
                font_size = 10.5, 
                line_spacing = 0, 
                space_before = 0,
                align = "justify", 
                keep_together = True, 
                keep_with_next = False, 
                page_break_before = False,
                widow_control = False, 
                set_bold = False, 
                set_italic = False, 
                set_underline = False, 
                set_all_caps = False,
                style_name = "Normal"):
    paragraph = document.add_paragraph(content)
    if WD_STYLE_TYPE.PARAGRAPH != 1:
        paragraph.style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    font = paragraph.style.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = set_bold
    font.italic = set_italic
    font.all_caps = set_all_caps
    font.underline = set_underline
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = alignment_dict.get(align.lower())
    paragraph_format.space_before = Pt(space_before)
    paragraph_format.space_after = Pt(space_after)
    # paragraph_format.line_spacing = line_spacing
    # paragraph_format.keep_together = keep_together
    # paragraph_format.keep_with_next = keep_with_next
    # paragraph_format.page_break_before = page_break_before
    # paragraph_format.widow_control = widow_control

def add_cell_content(cell,
                     index = 0,
                     content = "",
                     font_name = font_type,
                     font_size = 10.5,
                     align = "justify",
                     space_after = 0,
                     space_before = 0,
                     set_bold = False, 
                     set_italic = False, 
                     set_underline = False, 
                     set_all_caps = False,
                     style_name = "Normal"):
    cell.paragraphs[index].space_before = Pt(space_before)
    cell.paragraphs[index].space_after = Pt(space_after)
    cell.paragraphs[index].alignment = alignment_dict.get(align.lower())
    run = cell.paragraphs[index].add_run(content)
    if WD_STYLE_TYPE.PARAGRAPH != 1:
        run.style = cell.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    font = run.font
    font.name = font_name
    font.size = Pt(font_size)
    font.bold = set_bold
    font.italic = set_italic
    font.all_caps = set_all_caps
    font.underline = set_underline
    run.alignment = alignment_dict.get(align.lower())

############################################################################################################
#########
add_content(document = doc_file,
            content = "「実船計測及び海象予測に基づいた荷重／強度評価技術に関する研究」\n対象船のデータ取得状況",
            align = "Center",
            style_name = "Normal")
#########

#########
add_content(document = doc_file,
            content = date_time + "\n" + company_name,
            align = "Right",
            style_name = "Normal")
#########

#########
add_content(document = doc_file,
            content = "1. 構造モニタリングデータ取得状況",
            align = "Left",
            style_name = "Normal")
#########

#########
first_table = doc_file.add_table(rows = 1,
                                 cols = num_of_col)
first_table.style = 'Table Grid'
first_table.alignment = alignment_dict["center"]
####
title_row = first_table.rows[0].cells
add_cell_content(cell = title_row[0], 
                 index = 0, 
                 content = "No.",
                 align = "center",
                 style_name = "Normal")
title_row[0].width = Mm(10)
title_row[0].height = Mm(6.5)
add_cell_content(cell = title_row[1], 
                 index = 0, 
                 content = "データ名",
                 align = "center",
                 style_name = "Normal")
title_row[1].width = Mm(56)
title_row[1].height = Mm(6.5)
####
for i in range(2, len(title_row)):
    add_cell_content(cell = title_row[i], 
                     index = 0, 
                     content = col_titles[i-2],
                     align = "center",
                     style_name = "Normal")
####
for index, data in enumerate(row_data):
    data_row = first_table.add_row().cells
    add_cell_content(cell = data_row[0], 
                     index = 0, 
                     content = str(index + 1),
                     align = "center",
                     style_name = "Normal")
    data_row[0].width = Mm(10)
    data_row[0].height = Mm(6.5)
    data_row[1].width = Mm(56)
    data_row[1].height = Mm(6.5)
    for i in range(len(data)):
        data_row[i + 1].width = Mm(37.5)
        data_row[i + 1].height = Mm(6.5)
        if i == 0:
            add_cell_content(cell = data_row[i + 1], 
                             index = 0, 
                             content = data[i],
                             align = "left",
                             style_name = "Normal")
        else:
            add_cell_content(cell = data_row[i + 1], 
                             index = 0, 
                             content = data[i],
                             align = "center",
                             style_name = "Normal")
############################################################################################################

#########
add_content(document = doc_file,
            content = "\n以下に{0} の{1}分の取得データを示す。\n".format(getting_data[0], getting_data[1]),
            align = "Left",
            style_name = "Normal")
#########

#########
add_content(document = doc_file,
            content = "2. 航路",
            align = "Left",
            style_name = "Normal")
#########

#########
add_content(document = doc_file,
            content = " {0}の就航後の航路を図1 に示す。図中の色の淡い方から濃い方に向けて本船が航行している。".format(getting_data[0], getting_data[1]),
            align = "Left",
            style_name = "Normal")
#########

#########
add_picture(document = doc_file,
            path = os.getcwd() + "/" + "images/Map.PNG",
            align = "Center",
            pic_width = pic_size[0][0],
            pic_height = pic_size[0][1])
#########

#########
add_content(document = doc_file,
            content = "図1 航路",
            align = "Center",
            style_name = "Normal")
#########

doc_file.add_page_break()

#########
add_content(document = doc_file,
            content = "3. 応力",
            align = "Left",
            style_name = "Normal")
#########

#########
add_content(document = doc_file,
            content = " 計測した応力時系列を図２に示す。図中の青線が{0}分平均値、誤差範囲として{0}分標準偏差をプロットしている。\n".format(getting_data[2]),
            align = "Left",
            style_name = "Normal")
#########

for index, items in enumerate(picture_2_group):
    #########
    add_picture(document = doc_file,
                path = items[1],
                align = "Center")
    #########

    #########
    add_content(document = doc_file,
                content = str(index + 1) + ".　" + items[0] + "\n",
                align = "Center",
                style_name = "Normal")
    #########

#########
add_content(document = doc_file,
            content = "図2 計測応力時系列",
            align = "Center",
            style_name = "Normal")
#########

doc_file.add_page_break()

#########
add_content(document = doc_file,
            content = "4. 加速度",
            align = "Left",
            style_name = "Normal")
#########

#########
add_content(document = doc_file,
            content = " 計測した加速度時系列を図3 に示す。図中の青線が{0}平均値、誤差範囲として{0}標準偏差をプロットしている。\n".format(getting_data[2]),
            align = "Left",
            style_name = "Normal")
#########

for index, items in enumerate(picture_3_group):
    #########
    add_picture(document = doc_file,
                path = items[1],
                align = "Center")
    #########

    #########
    add_content(document = doc_file,
                content = str(index + 1 + len(picture_2_group)) + ".　" + items[0] + "\n",
                align = "Center",
                style_name = "Normal")
    #########

#########
add_content(document = doc_file,
            content = "図3 計測加速度時系列",
            align = "Center",
            style_name = "Normal")
#########

doc_file.add_page_break()

#########
add_content(document = doc_file,
            content = "5. 運動",
            align = "Left",
            style_name = "Normal")
#########

#########
add_content(document = doc_file,
            content = " ジャイロで計測したRoll とPitch の{0}最大値をプロットしたものを図４に示す。\n".format(getting_data[2]),
            align = "Left",
            style_name = "Normal")
########

for index, items in enumerate(picture_4_group):
    #########
    add_picture(document = doc_file,
                path = items[1],
                align = "Center")
    #########

    #########
    add_content(document = doc_file,
                content = str(index + 1 + len(picture_2_group) + len(picture_3_group)) + ".　" + items[0] + "\n",
                align = "Center",
                style_name = "Normal")
    #########

#########
add_content(document = doc_file,
            content = "図4 運動時系列",
            align = "Center",
            style_name = "Normal")
#########

doc_file.add_page_break()

#########
add_content(document = doc_file,
            content = "6. 海象予報値",
            align = "Left",
            style_name = "Normal")
#########

#########
add_content(document = doc_file,
            content = " {0} の航路データから{1} より得た海象予報値のうち、有義波高と波周期を図5 に示す。\n".format(col_titles[1], getting_data[3]),
            align = "Left",
            style_name = "Normal")
########

for index, items in enumerate(picture_5_group):
    #########
    add_picture(document = doc_file,
                path = items[1],
                align = "Center")
    #########

    #########
    add_content(document = doc_file,
                content = items[0] + "\n",
                align = "Center",
                style_name = "Normal")
    #########

#########
add_content(document = doc_file,
            content = "図5 海象予報値",
            align = "Center",
            style_name = "Normal")
#########

#########
add_content(document = doc_file,
            content = "\n\n\n以上",
            align = "right",
            style_name = "Normal")
#########

doc_file.save(file_path)
convert(file_path, file_path_pdf)